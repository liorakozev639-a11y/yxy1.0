from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\杨星宇\Documents\Codex\2026-08-01\an-zhu")
SOURCE = ROOT / "outputs" / "free-time-agent-product-plan.md"
OUTPUT = ROOT / "outputs" / "空闲时间规划Agent产品计划书.docx"

BLUE = "2E74B5"
NAVY = "0B2545"
DARK_BLUE = "1F4D78"
GRAY = "59636E"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_rich_text(paragraph, text, size=10.5, color=None):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        code = part.startswith("`") and part.endswith("`")
        clean = part[2:-2] if bold else part[1:-1] if code else part
        run = paragraph.add_run(clean)
        set_run_font(run, size=size, bold=bold, color=(DARK_BLUE if code else color))
        if code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def set_paragraph(p, before=0, after=6, line=1.1, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text, style=None, before=0, after=6, line=1.1, color=None, size=10.5, bold=False):
    p = doc.add_paragraph(style=style)
    set_paragraph(p, before=before, after=after, line=line)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph(p, after=4, line=1.167)
    add_rich_text(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph(p, after=4, line=1.167)
    add_rich_text(p, text)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_rich_text(p, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    widths = [1440, 2520, 2520, 2880] if cols == 4 else [3120, 6240] if cols == 2 else [2000] * cols
    if sum(widths) != 9360:
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, val in enumerate(rows[0]):
        set_cell_shading(hdr.cells[i], LIGHT_GRAY)
        hdr.cells[i].text = ""
        p = hdr.cells[i].paragraphs[0]
        set_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(val)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_paragraph(p, after=0, line=1.05)
            add_rich_text(p, val, size=9.3)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_paragraph(p, before=2, after=2, line=1.1)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    add_rich_text(p, text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "等线"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, color, before, after in [(1,16,BLUE,16,8),(2,13,BLUE,12,6),(3,12,DARK_BLUE,8,4)]:
        st = doc.styles[f"Heading {level}"]
        st.font.name = "等线"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Bullet 2", "List Number"]:
        st = doc.styles[name]
        st.font.name = "等线"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167


def add_header_footer(doc):
    sec = doc.sections[0]
    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p, after=0, line=1.0)
    r = p.add_run("空闲时间规划 Agent | 产品计划书")
    set_run_font(r, size=8.5, color=GRAY)
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=0, line=1.0)
    r = p.add_run("内部产品方案 | 2026-08-01")
    set_run_font(r, size=8.5, color=GRAY)


def parse_markdown(doc, text):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        m = re.match(r"^(#{2,4})\s+(.*)$", line)
        if m:
            add_heading(doc, m.group(2).strip(), len(m.group(1)) - 1)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                vals = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(set(v) <= set("-: ") for v in vals):
                    rows.append(vals)
                i += 1
            if rows:
                add_table(doc, rows)
            continue
        if re.match(r"^[-*]\s+", line):
            add_bullet(doc, re.sub(r"^[-*]\s+", "", line))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            add_number(doc, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue
        if line.startswith(">"):
            add_callout(doc, "关键结论", line[1:].strip())
            i += 1
            continue
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{2,4})\s|^[-*]\s|^\d+\.\s|^\||^>", lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        set_paragraph(p, after=6, line=1.1)
        add_rich_text(p, " ".join(para_lines))


def main():
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    set_paragraph(p, before=70, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("空闲时间规划 Agent")
    set_run_font(r, size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_paragraph(p, after=20, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("产品计划书")
    set_run_font(r, size=20, bold=True, color=BLUE)
    p = doc.add_paragraph()
    set_paragraph(p, after=28, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("把突然拥有的空闲时间，变成一份真正能执行的安排")
    set_run_font(r, size=12.5, color=GRAY)

    add_callout(doc, "文档定位", "面向产品、设计、研发和运营团队的 MVP 方案。当前版本聚焦“偏好采集 - 任务推荐 - 排程 - 执行 - 反馈”的闭环，并明确首版的边界、兜底和隐私要求。")

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [2160, 7200])
    pairs = [("版本", "V1.0 产品方案"), ("日期", "2026-08-01"), ("目标用户", "在校学生；长期工作繁忙、突然获得休闲时间后不知道如何规划的职场人"), ("交付形态", "网页展示、PDF 下载、单次授权邮件发送")]
    for idx, (k, v) in enumerate(pairs):
        for j, val in enumerate((k, v)):
            c = meta.cell(idx, j)
            c.text = ""
            if j == 0:
                set_cell_shading(c, LIGHT_GRAY)
            p = c.paragraphs[0]
            set_paragraph(p, after=0, line=1.0)
            r = p.add_run(val)
            set_run_font(r, size=9.5, bold=(j == 0), color=NAVY if j == 0 else None)
    set_table_geometry(meta, [2160, 7200])
    doc.add_page_break()

    add_heading(doc, "目录", 1)
    toc = [
        "1. 产品概述", "2. 用户与使用场景", "3. 兴趣方向与前置条件", "4. 用户流程",
        "5. 问卷规则", "6. 系统设计", "7. 交付与隐私", "8. 异常处理与质量控制",
        "9. 验收标准", "10. 阶段规划", "11. 成功指标"
    ]
    for item in toc:
        add_para(doc, item, after=4, color=NAVY, size=10.5)
    add_callout(doc, "阅读提示", "本文是第一版产品计划书，功能优先级以 MVP 闭环为中心。地图、活动、商户、日历、通知和邮件均通过适配器接入，外部服务不可用时必须保留可用的通用任务与网页计划。")

    parse_markdown(doc, text)
    doc.core_properties.title = "空闲时间规划 Agent 产品计划书"
    doc.core_properties.subject = "空闲时间规划 Agent MVP 产品方案"
    doc.core_properties.author = "产品团队"
    doc.core_properties.comments = "Generated from the confirmed product plan."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
